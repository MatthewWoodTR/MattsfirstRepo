#!/usr/bin/env python3
import sys
import os
import re
import subprocess
from pathlib import Path
from typing import Dict, List

from pdf2image import convert_from_path
from docx import Document
from docx.shared import Inches

PDF = Path(sys.argv[1]) if len(sys.argv) > 1 else Path(os.environ.get('PDF_PATH', 'Proposed enhancements for Anchin.pdf'))
PRD_DIR = Path(sys.argv[2]) if len(sys.argv) > 2 else Path(os.environ.get('PRD_DIR', 'PRDs/2025-11'))
ASSETS_ROOT = PRD_DIR / 'assets'
ALL_PAGES_DIR = ASSETS_ROOT / '_all_pages'

# Keyword groups for epics
DIVISION_EPICS = {
    '4233645', '4233303', '4379342', '4233299', '4233646', '4233648', '4235716'
}
NOTES_EPICS = {
    '4205843', '4237582', '4237578'
}
DRILL_EPICS = {'4204113'}
CONSOLIDATION_EPICS = {
    '4237546', '4233330', '4206592', '4233327', '4233310', '4379363', '4442269'
}

KEYWORDS = {
    'division': 'DIVISION',
    'division set': 'DIVISION',
    'note': 'NOTES',
    'notes': 'NOTES',
    'drill': 'DRILL',
    'formula': 'DRILL',
    'excel': 'DRILL',
    'word': 'DRILL',
    'consolidation': 'CONSOLIDATION',
    'child engagement': 'CONSOLIDATION',
    'finalize': 'CONSOLIDATION',
    'roll forward': 'CONSOLIDATION',
    'spreadsheet': 'CONSOLIDATION',
}

CATEGORY_TO_EPICS = {
    'DIVISION': DIVISION_EPICS,
    'NOTES': NOTES_EPICS,
    'DRILL': DRILL_EPICS,
    'CONSOLIDATION': CONSOLIDATION_EPICS,
}


def ensure_dirs():
    ALL_PAGES_DIR.mkdir(parents=True, exist_ok=True)


def pdftotext_page(pdf_path: Path, page_num: int) -> str:
    # Extract text for a single page using poppler-utils
    tmp_txt = ALL_PAGES_DIR / f"page-{page_num}.txt"
    try:
        subprocess.run([
            'pdftotext', '-f', str(page_num), '-l', str(page_num), '-layout', str(pdf_path), str(tmp_txt)
        ], check=True)
        return tmp_txt.read_text(encoding='utf-8', errors='ignore')
    except Exception:
        return ''


def categorize_page(text: str) -> str:
    t = text.lower()
    for kw, cat in KEYWORDS.items():
        if kw in t:
            return cat
    # fallback unknown
    return 'UNKNOWN'


def convert_pdf_to_images(pdf_path: Path) -> List[Path]:
    # 2x resolution for clarity
    images = convert_from_path(str(pdf_path), dpi=200)
    saved = []
    for i, img in enumerate(images, start=1):
        out = ALL_PAGES_DIR / f"page-{i}.png"
        img.save(out, 'PNG')
        saved.append(out)
    return saved


def find_epic_id_from_filename(name: str) -> str:
    m = re.match(r"^(\d{7})\s-\s", name)
    return m.group(1) if m else ''


def append_images_to_docx(docx_path: Path, images: List[Path]):
    doc = Document(str(docx_path))
    # Add heading
    doc.add_heading('Appendix: Quick prototype', level=1)
    for img in images:
        # Scale to page width (~6.5 inches typical content width). Adjust if needed.
        doc.add_picture(str(img), width=Inches(6.5))
    doc.save(str(docx_path))


def main():
    assert PDF.exists(), f"PDF not found: {PDF}"
    assert PRD_DIR.exists(), f"PRD dir not found: {PRD_DIR}"

    ensure_dirs()

    # Convert PDF to page images once
    page_images = convert_pdf_to_images(PDF)

    # Extract text per page to categorize
    page_categories: Dict[int, str] = {}
    for i in range(1, len(page_images) + 1):
        txt = pdftotext_page(PDF, i)
        cat = categorize_page(txt)
        page_categories[i] = cat

    # Map categories to page indices
    cat_to_pages: Dict[str, List[int]] = {}
    for idx, cat in page_categories.items():
        cat_to_pages.setdefault(cat, []).append(idx)

    # Process each RTF: convert to DOCX, then append images for its epic/group
    rtf_files = sorted([p for p in PRD_DIR.glob('*.rtf')])
    for rtf in rtf_files:
        base_name = rtf.name
        epic_id = find_epic_id_from_filename(base_name)
        if not epic_id:
            print(f"Skip (no epic id): {base_name}")
            continue

        # Convert RTF -> DOCX with pandoc
        docx_out = rtf.with_suffix('.docx')
        subprocess.run(['pandoc', str(rtf), '-o', str(docx_out)], check=True)

        # Determine category by epic id
        if epic_id in DIVISION_EPICS:
            cat = 'DIVISION'
        elif epic_id in NOTES_EPICS:
            cat = 'NOTES'
        elif epic_id in DRILL_EPICS:
            cat = 'DRILL'
        elif epic_id in CONSOLIDATION_EPICS:
            cat = 'CONSOLIDATION'
        else:
            cat = 'UNKNOWN'

        # Collect images
        page_idxs = cat_to_pages.get(cat, [])
        # Fallback: if no pages match, include the first page as a placeholder
        if not page_idxs and len(page_images) > 0:
            page_idxs = [1]

        # Copy selected images under per-epic assets and append to docx
        per_epic_dir = ASSETS_ROOT / epic_id
        per_epic_dir.mkdir(parents=True, exist_ok=True)
        selected_imgs: List[Path] = []
        for idx in page_idxs:
            src = ALL_PAGES_DIR / f"page-{idx}.png"
            dst = per_epic_dir / src.name
            if not dst.exists():
                # copy by writing bytes
                dst.write_bytes(src.read_bytes())
            selected_imgs.append(dst)

        append_images_to_docx(docx_out, selected_imgs)
        print(f"Updated {docx_out} with {len(selected_imgs)} image(s) for epic {epic_id} ({cat})")

if __name__ == '__main__':
    main()
