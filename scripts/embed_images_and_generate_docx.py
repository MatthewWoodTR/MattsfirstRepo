import os
import re
from pathlib import Path
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches

REPO_ROOT = Path(__file__).resolve().parents[1]
PRD_DIR = REPO_ROOT / 'PRDs' / '2025-11'
PDF_PATH = REPO_ROOT / 'Proposed enhancements for Anchin.pdf'
ASSETS_ROOT = PRD_DIR / 'assets'

CATEGORY_KEYWORDS = {
    'consolidation': ['consolidation', 'parent', 'child', 'roll forward', 'finalize', 'spreadsheet import'],
    'division': ['division', 'division set', 'divisional', 'tax product export'],
    'notes': ['note', 'notes', 'dialog', 'format'],
    'drill': ['drill-down', 'drilldown', 'formula', 'excel', 'word']
}

# Simple filename-to-category heuristic
def infer_categories_from_filename(name: str):
    lower = name.lower()
    cats = set()
    for cat, kws in CATEGORY_KEYWORDS.items():
        if any(kw in lower for kw in kws):
            cats.add(cat)
    return cats or {'consolidation', 'division', 'notes', 'drill'}

# Extract PDF pages to images and collect per-page text for routing
class PdfExtractor:
    def __init__(self, pdf_path: Path):
        self.doc = fitz.open(pdf_path)
        self.pages = []
        for i, page in enumerate(self.doc):
            text = page.get_text("text").strip()
            self.pages.append({
                'index': i,
                'text': text,
            })

    def export_pages(self, out_dir: Path):
        out_dir.mkdir(parents=True, exist_ok=True)
        exported = []
        for p in self.pages:
            page = self.doc[p['index']]
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x scale
            out_file = out_dir / f"page-{p['index'] + 1}.png"
            pix.save(out_file)
            exported.append({**p, 'image': out_file})
        return exported

# Match pages to PRDs by category and title heuristics
def pages_for_prd(prd_filename: str, pages_meta):
    name = prd_filename.lower()
    cats = infer_categories_from_filename(name)
    matched = []
    for p in pages_meta:
        t = (p['text'] or '').lower()
        # Category filter by text
        if 'consolidation' in cats and any(k in t for k in ['consolidation', 'parent', 'child']):
            matched.append(p)
            continue
        if 'division' in cats and 'division' in t:
            matched.append(p)
            continue
        if 'notes' in cats and 'note' in t:
            matched.append(p)
            continue
        if 'drill' in cats and any(k in t for k in ['drill', 'formula', 'excel', 'word']):
            matched.append(p)
            continue
    # Fallback: if nothing matched, include first page
    if not matched and pages_meta:
        matched = [pages_meta[0]]
    # De-duplicate while preserving order
    seen = set()
    unique = []
    for m in matched:
        idx = m['index']
        if idx not in seen:
            unique.append(m)
            seen.add(idx)
    return unique


def convert_rtf_to_plaintext(rtf_path: Path) -> str:
    # Minimal RTF cleanup — strips common control words; retains readable text
    data = rtf_path.read_text(encoding='utf-8', errors='ignore')
    # Remove RTF header and groups
    data = re.sub(r'\\pard', '\n', data)
    data = re.sub(r'\\par', '\n', data)
    data = re.sub(r'\\tab', '    ', data)
    data = re.sub(r'\\[^\\ ]+ ?', '', data)  # remove control words
    data = re.sub(r'[{}]', '', data)  # remove group braces
    # Collapse excessive newlines
    data = re.sub(r'\n\s*\n+', '\n\n', data)
    return data.strip()


def build_docx_from_rtf(rtf_path: Path, pages_meta, images_root: Path):
    # Epic ID from filename prefix
    epic_id = rtf_path.name.split(' ')[0]
    prd_title = rtf_path.stem

    doc = Document()
    doc.add_heading(prd_title, level=1)

    # Pull existing content as plain text
    text = convert_rtf_to_plaintext(rtf_path)
    if text:
        for para in text.split('\n\n'):
            doc.add_paragraph(para)

    # Quick prototype section
    doc.add_heading('Appendix: Quick prototype', level=2)

    # Export PDF pages to images under epic-specific dir
    epic_img_dir = images_root / epic_id
    epic_img_dir.mkdir(parents=True, exist_ok=True)

    matched_pages = pages_for_prd(rtf_path.name, pages_meta)

    for meta in matched_pages:
        # Ensure page image exists in epic dir
        source_img = meta['image']
        # Copy to epic dir (linking by re-saving)
        with open(source_img, 'rb') as fsrc:
            img_bytes = fsrc.read()
        target_img = epic_img_dir / source_img.name
        with open(target_img, 'wb') as fdst:
            fdst.write(img_bytes)

        doc.add_paragraph(f"Figure: PDF page {meta['index'] + 1}")
        try:
            doc.add_picture(str(target_img), width=Inches(6.5))
        except Exception as e:
            doc.add_paragraph(f"[Image insertion failed: {e}]")

    # ADO link appendix
    doc.add_heading('Appendix: Links', level=2)
    ado_id = epic_id
    ado_url = f"https://dev.azure.com/tr-tax/TaxProf/_workitems/edit/{ado_id}"
    p = doc.add_paragraph()
    run = p.add_run("ADO Epic Link: ")
    run = p.add_run(ado_url)

    # Save DOCX alongside RTF
    out_docx = rtf_path.with_suffix('.docx')
    doc.save(out_docx)


def main():
    if not PDF_PATH.exists():
        print(f"PDF not found at {PDF_PATH}")
        return

    extractor = PdfExtractor(PDF_PATH)

    # Export all pages once to a shared cache under assets/_all_pages
    shared_pages_dir = ASSETS_ROOT / '_all_pages'
    pages_meta = extractor.export_pages(shared_pages_dir)

    # For each PRD RTF, create a DOCX with embedded images
    for rtf in sorted(PRD_DIR.glob('*.rtf')):
        build_docx_from_rtf(rtf, pages_meta, ASSETS_ROOT)

if __name__ == '__main__':
    main()
