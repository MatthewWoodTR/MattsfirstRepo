#!/usr/bin/env python3
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

ROOT = Path('PRDs/2025-11')
HEADERS = [
    '1. Customer Problem',
    '2. Customer Research',
    '3. Our Solution',
    '4. Product Metrics',
    'Appendix: Additional Links',
    'Appendix: Quick prototype'
]

ENTERPRISE_SENTENCE = (
    'This capability was requested as feedback from an enterprise-level accounting firm, '
    'reflecting needs observed in large multi-entity audit workflows.'
)

COMPETITIVE_SENTENCE = (
    'We are also building this to achieve competitive parity with Wolters Kluwer ProSystem fx Engagement, '
    'which offers similar functionality.'
)

def _para_text(elem):
    return ''.join(t.text for t in elem.iter(qn('w:t'))).strip()


def insert_paragraph_after(paragraph: Paragraph, text: str = '', style: str | None = None) -> Paragraph:
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style:
        new_para.style = style
    return new_para


def paragraph_is_header(p: Paragraph):
    txt = (p.text or '').strip()
    return any(txt.startswith(h) for h in HEADERS)


def ensure_header_style(p: Paragraph):
    # Make entire paragraph bold
    if not p.runs:
        p.add_run('')
    for r in p.runs:
        r.bold = True
    # Ensure a blank line before and after (idempotent)
    prev = p._p.getprevious()
    if prev is None or _para_text(prev) != '':
        new_prev = OxmlElement('w:p')
        p._p.addprevious(new_prev)
    nxt = p._p.getnext()
    if nxt is None or _para_text(nxt) != '':
        new_next = OxmlElement('w:p')
        p._p.addnext(new_next)
    fmt = p.paragraph_format
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(6)


def add_research_sentences(doc: Document, is_div_or_cons: bool):
    # Add sentences immediately under '2. Customer Research' header; avoid duplicates
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().lower().startswith('2. customer research'):
            window_text = '\n'.join(q.text.strip() for q in doc.paragraphs[i+1:i+6])
            need_enterprise = ENTERPRISE_SENTENCE not in window_text
            need_competitive = is_div_or_cons and (COMPETITIVE_SENTENCE not in window_text)
            if not (need_enterprise or need_competitive):
                return True
            # ensure one blank line after header
            after = insert_paragraph_after(p, '')
            if need_enterprise:
                after = insert_paragraph_after(after, ENTERPRISE_SENTENCE)
            if need_competitive:
                after = insert_paragraph_after(after, COMPETITIVE_SENTENCE)
            return True
    return False


def process_docx(path: Path):
    print(f'Processing {path}')
    doc = Document(str(path))

    name_lower = path.name.lower()
    is_div_or_cons = ('division' in name_lower) or ('consolidation' in name_lower)

    for p in doc.paragraphs:
        if paragraph_is_header(p):
            ensure_header_style(p)

    add_research_sentences(doc, is_div_or_cons)

    doc.save(str(path))


def main():
    if not ROOT.exists():
        print('No PRDs directory found; skipping.')
        return
    for f in ROOT.glob('*.docx'):
        process_docx(f)

if __name__ == '__main__':
    main()
