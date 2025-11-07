import os
import re
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_BREAK

PRDS_DIR = "PRDs/2025-11"
APPENDIX_HEADER_PATTERNS = [
    re.compile(r"^appendix\s*:\s*quick\s*prototype\s*$", re.IGNORECASE),
    re.compile(r"^appendix\s*[-:–—]\s*quick\s*prototype\s*$", re.IGNORECASE),
]

REMOVE_PHRASES = [
    re.compile(r"\bfrom the epic\b", re.IGNORECASE),
    re.compile(r"\bfrom the feature\b", re.IGNORECASE),
]

TOP_LEVEL_HEADERS = [
    re.compile(r"^1\.\s*customer\s*problem", re.IGNORECASE),
    re.compile(r"^2\.\s*customer\s*research", re.IGNORECASE),
    re.compile(r"^3\.\s*our\s*solution", re.IGNORECASE),
    re.compile(r"^4\.\s*product\s*metrics", re.IGNORECASE),
    re.compile(r"^appendix", re.IGNORECASE),
]

def paragraph_has_image(p):
    # Detect embedded pictures by scanning runs for drawing elements
    for run in p.runs:
        r = run._r
        if r.xpath('.//w:drawing') or r.xpath('.//w:pict'):
            return True
    return False


def is_header_paragraph(p):
    text = (p.text or "").strip()
    for pat in TOP_LEVEL_HEADERS:
        if pat.match(text):
            return True
    return False


def should_start_appendix_deletion(p):
    txt = (p.text or "").strip()
    for pat in APPENDIX_HEADER_PATTERNS:
        if pat.match(txt):
            return True
    return False


def scrub_phrases_in_paragraph(p):
    # Replace phrases in paragraph text while preserving runs where possible
    if not p.text:
        return
    new_text = p.text
    for rx in REMOVE_PHRASES:
        new_text = rx.sub("", new_text)
    # Clean up double spaces and stray punctuation
    new_text = re.sub(r"\s{2,}", " ", new_text)
    new_text = re.sub(r"\s+([.,;:])", r"\1", new_text)
    if new_text != p.text:
        # Replace all runs with single run containing new_text to ensure consistency
        for r in list(p.runs):
            r.clear()
        p.runs[0].add_text(new_text) if p.runs else p.add_run(new_text)


def remove_appendix_and_images(doc: Document):
    paras = list(doc.paragraphs)
    i = 0
    while i < len(paras):
        p = paras[i]
        if should_start_appendix_deletion(p):
            # Delete this paragraph and following content until next top-level header or end
            j = i
            while j < len(paras):
                q = paras[j]
                if j != i and is_header_paragraph(q):
                    break
                # remove paragraph q
                q._element.getparent().remove(q._element)
                j += 1
            # Refresh paragraphs snapshot and continue from i (same index now points to next element)
            paras = list(doc.paragraphs)
            continue
        i += 1

    # Also remove any images elsewhere in the document
    for p in list(doc.paragraphs):
        if paragraph_has_image(p):
            # Remove the image-containing paragraph
            p._element.getparent().remove(p._element)


def process_docx(path):
    doc = Document(path)
    remove_appendix_and_images(doc)
    for p in list(doc.paragraphs):
        scrub_phrases_in_paragraph(p)
    doc.save(path)


def main():
    for root, _, files in os.walk(PRDS_DIR):
        for name in files:
            if name.lower().endswith('.docx'):
                fpath = os.path.join(root, name)
                print(f"Cleaning {fpath}")
                process_docx(fpath)

if __name__ == '__main__':
    main()
